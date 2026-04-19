"""공식 제안서 문서 생성기.

DOCX(Word) 형식으로 출력합니다.
HWP는 오픈소스 라이브러리 미지원으로 DOCX를 제공하며,
한글(HWP) 프로그램에서 DOCX를 직접 열 수 있습니다.
"""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DOCS_DIR = Path("./data/documents")
DOCS_DIR.mkdir(parents=True, exist_ok=True)


def _set_cell_bg(cell, hex_color: str) -> None:
    """표 셀 배경색 설정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)  # blue-700
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)


def _divider(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "93C5FD")  # blue-300
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(
    proposal: dict,
    classification: str,
    analysis: dict | None,
    session_id: str,
) -> Path:
    """공식 제안서 DOCX 생성 후 저장 경로 반환."""
    doc = Document()

    # ── 페이지 여백 ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # ── 헤더 타이틀 ─────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(proposal.get("title", "제안서"))
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph()

    # ── 메타 정보 표 ────────────────────────────────────────────────────────────
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_data = [
        ("문서 유형", classification),
        ("담당 부처", proposal.get("responsible_dept", "-")),
        ("제출일", datetime.now().strftime("%Y년 %m월 %d일")),
        ("접수번호", session_id[:8].upper()),
    ]
    for i, (label, value) in enumerate(meta_data):
        label_cell = meta_table.rows[i].cells[0]
        value_cell = meta_table.rows[i].cells[1]
        _set_cell_bg(label_cell, "DBEAFE")
        label_cell.text = label
        label_cell.paragraphs[0].runs[0].bold = True
        label_cell.paragraphs[0].runs[0].font.size = Pt(10)
        value_cell.text = value
        value_cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    _divider(doc)

    # ── 1. 제안 배경 ────────────────────────────────────────────────────────────
    _heading(doc, "1. 제안 배경")
    _body(doc, proposal.get("background", "-"))
    _divider(doc)

    # ── 2. 주요 내용 ────────────────────────────────────────────────────────────
    _heading(doc, "2. 주요 요청 사항")
    core = proposal.get("core_requests", "-")
    for line in core.split("\n"):
        line = line.strip()
        if line:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(line.lstrip("•-·").strip())
            run.font.size = Pt(10.5)
    _divider(doc)

    # ── 3. 기대 효과 ────────────────────────────────────────────────────────────
    _heading(doc, "3. 기대 효과")
    effects = proposal.get("expected_effects", "-")
    for line in effects.split("\n"):
        line = line.strip()
        if line:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(line.lstrip("•-·").strip())
            run.font.size = Pt(10.5)
    _divider(doc)

    # ── 4. 관련 법령 ────────────────────────────────────────────────────────────
    _heading(doc, "4. 관련 법령")
    laws = proposal.get("related_laws", [])
    if laws:
        for law in laws:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(law if isinstance(law, str) else law.get("title", str(law)))
            run.font.size = Pt(10.5)
    else:
        _body(doc, "관련 법령 없음")
    _divider(doc)

    # ── 5. AI 분석 결과 ─────────────────────────────────────────────────────────
    if analysis:
        _heading(doc, "5. AI 분석 결과")

        analysis_table = doc.add_table(rows=3, cols=2)
        analysis_table.style = "Table Grid"
        feasibility = analysis.get("feasibility_score", 0)
        pass_prob = analysis.get("pass_probability", 0)
        duration = analysis.get("expected_duration_days", 0)
        metrics = [
            ("실현 가능성", f"{feasibility * 100:.0f}%"),
            ("처리 통과 예상 확률", f"{pass_prob * 100:.0f}%"),
            ("예상 처리 기간", f"약 {duration}일"),
        ]
        for i, (label, value) in enumerate(metrics):
            lc = analysis_table.rows[i].cells[0]
            vc = analysis_table.rows[i].cells[1]
            _set_cell_bg(lc, "F0FDF4")
            lc.text = label
            lc.paragraphs[0].runs[0].bold = True
            lc.paragraphs[0].runs[0].font.size = Pt(10)
            vc.text = value
            vc.paragraphs[0].runs[0].font.size = Pt(10)
            vc.paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

    # ── 푸터 ────────────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"본 문서는 JUT_AI신문고 시스템에 의해 자동 생성되었습니다.\n"
        f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')} | 접수번호: {session_id[:8].upper()}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # ── 저장 ────────────────────────────────────────────────────────────────────
    safe_title = proposal.get("title", "제안서")[:30].replace("/", "_").replace("\\", "_")
    filename = f"{session_id[:8]}_{safe_title}.docx"
    save_path = DOCS_DIR / filename
    doc.save(str(save_path))
    return save_path


def docx_to_bytes(path: Path) -> bytes:
    return path.read_bytes()
